"""
Pitboss v4 - DOCX Writer

Writes Document objects back to .docx files with formatting preserved.
This is the inverse of ingest/docx_reader.py.
"""

from __future__ import annotations
from pathlib import Path
import logging

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.document import (
    Document,
    Paragraph,
    Heading,
    HeadingLevel,
    List,
    ListItem,
    ListType,
    Table,
    TextRun,
)


logger = logging.getLogger(__name__)


# Map highlight color names to Word highlight constants
HIGHLIGHT_MAP = {
    'yellow': WD_COLOR_INDEX.YELLOW,
    'bright green': WD_COLOR_INDEX.BRIGHT_GREEN,
    'green': WD_COLOR_INDEX.BRIGHT_GREEN,
    'turquoise': WD_COLOR_INDEX.TURQUOISE,
    'cyan': WD_COLOR_INDEX.TURQUOISE,
    'pink': WD_COLOR_INDEX.PINK,
    'magenta': WD_COLOR_INDEX.PINK,
    'blue': WD_COLOR_INDEX.BLUE,
    'red': WD_COLOR_INDEX.RED,
    'dark blue': WD_COLOR_INDEX.DARK_BLUE,
    'teal': WD_COLOR_INDEX.TEAL,
    'dark red': WD_COLOR_INDEX.DARK_RED,
    'dark yellow': WD_COLOR_INDEX.DARK_YELLOW,
    'gray': WD_COLOR_INDEX.GRAY_25,
    'gray 25': WD_COLOR_INDEX.GRAY_25,
    'gray 50': WD_COLOR_INDEX.GRAY_50,
}

# Map heading levels to Word styles
HEADING_STYLES = {
    HeadingLevel.H1: 'Heading 1',
    HeadingLevel.H2: 'Heading 2',
    HeadingLevel.H3: 'Heading 3',
    HeadingLevel.H4: 'Heading 4',
}


from typing import Optional


def _is_empty_paragraph(element) -> bool:
    """Check if element is an existing blank row (empty paragraph)."""
    return isinstance(element, Paragraph) and not element.text.strip()


def write_docx(
    document: Document,
    output_path: Path,
    blank_rows: Optional[str] = None,
) -> None:
    """
    Write a Document object to a .docx file.

    Preserves:
    - Paragraph structure
    - Heading levels (H1-H4)
    - Lists (ordered/unordered)
    - Tables
    - Inline formatting (bold, italic, underline, strikethrough)
    - Highlights
    - Hyperlinks

    Args:
        document: The Document to write
        output_path: Path for the output .docx file
        blank_rows: Blank row handling mode:
            - "required": Insert empty paragraphs between content blocks
              (headings, paragraphs, lists, tables). Idempotent - won't
              double-insert if source already has empty rows. Lists exempt
              (no empty rows between list items). Tables exempt from
              table-to-table insertion (corpus: 100% directly adjacent).
            - "none": Don't insert blank rows (preserve source as-is)
            - None: Preserve source as-is (same as "none")
    """
    doc = DocxDocument()
    elements = document.elements

    for i, element in enumerate(elements):
        # Check if we need to insert empty row BEFORE this element
        if blank_rows == "required" and i > 0:
            prev_element = elements[i - 1]
            needs_empty = (
                not _is_empty_paragraph(element) and  # Current isn't already empty
                not _is_empty_paragraph(prev_element) and  # Prev isn't empty
                # Exempt table-to-table adjacency (corpus: 100% directly adjacent)
                not (isinstance(element, Table) and isinstance(prev_element, Table))
            )
            if needs_empty:
                doc.add_paragraph()  # Empty paragraph

        # Write the actual element
        if isinstance(element, Heading):
            _write_heading(doc, element)
        elif isinstance(element, Paragraph):
            _write_paragraph(doc, element)
        elif isinstance(element, List):
            _write_list(doc, element)
        elif isinstance(element, Table):
            _write_table(doc, element)

    # Ensure parent directory exists
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc.save(str(output_path))
    logger.info(f"Wrote document to {output_path}")


def _write_heading(doc: DocxDocument, heading: Heading) -> None:
    """Write a heading element with preserved formatting."""
    style = HEADING_STYLES.get(heading.level, 'Heading 1')

    # Add heading with style
    para = doc.add_heading(level=heading.level.value)
    para.clear()  # Remove default run

    # Add formatted runs
    _add_runs_to_paragraph(para, heading.runs())

    # Apply preserved paragraph formatting
    _apply_paragraph_formatting(para, heading)


def _write_paragraph(doc: DocxDocument, paragraph: Paragraph) -> None:
    """Write a paragraph element with preserved formatting."""
    para = doc.add_paragraph()
    _add_runs_to_paragraph(para, paragraph.runs())

    # Apply preserved paragraph formatting
    _apply_paragraph_formatting(para, paragraph)


def _write_list(doc: DocxDocument, lst: List) -> None:
    """
    Write a list element with proper Word numbering XML.

    Creates actual Word list paragraphs with w:numPr elements that
    will be recognized when the document is read back.
    """
    # Ensure document has numbering definitions
    _ensure_numbering_definitions(doc)

    # Get numId based on list type (1=bullets, 2=numbers by our definition)
    num_id = 2 if lst.list_type == ListType.ORDERED else 1

    for item in lst.items:
        para = doc.add_paragraph()
        _add_runs_to_paragraph(para, item.runs())

        # Apply proper numbering via XML
        _apply_list_numbering(para, num_id, item.indent_level)

        # Apply preserved paragraph formatting
        _apply_paragraph_formatting(para, item)


def _ensure_numbering_definitions(doc: DocxDocument) -> None:
    """
    Ensure document has numbering definitions for bullets and numbers.

    Creates the numbering.xml part if it doesn't exist, and adds
    abstract numbering definitions for bullet (numId=1) and numbered (numId=2) lists.
    """
    # Check if document already has numbering part
    numbering_part = doc.part.numbering_part

    # If numbering part exists, assume definitions are present
    # (this handles documents that already have lists)
    if numbering_part is not None:
        return

    # Create numbering definitions
    # This is done by adding a list item which triggers python-docx to create
    # the numbering infrastructure, then we can reference it
    # Note: python-docx creates numbering automatically when styles are used
    # We just need to ensure our numId values are valid

    # For now, we'll create the numbering XML directly
    _create_minimal_numbering(doc)


def _create_minimal_numbering(doc: DocxDocument) -> None:
    """
    Create minimal numbering definitions for bullets and numbers.

    This creates the numbering.xml part with two abstract number definitions:
    - abstractNumId="0" for bullets (numId=1)
    - abstractNumId="1" for numbers (numId=2)
    """
    from docx.parts.numbering import NumberingPart
    from docx.oxml.numbering import CT_Numbering

    # Create numbering element
    numbering_xml = """
    <w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <w:abstractNum w:abstractNumId="0">
            <w:lvl w:ilvl="0">
                <w:start w:val="1"/>
                <w:numFmt w:val="bullet"/>
                <w:lvlText w:val="●"/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="720" w:hanging="360"/>
                </w:pPr>
            </w:lvl>
        </w:abstractNum>
        <w:abstractNum w:abstractNumId="1">
            <w:lvl w:ilvl="0">
                <w:start w:val="1"/>
                <w:numFmt w:val="decimal"/>
                <w:lvlText w:val="%1."/>
                <w:lvlJc w:val="left"/>
                <w:pPr>
                    <w:ind w:left="720" w:hanging="360"/>
                </w:pPr>
            </w:lvl>
        </w:abstractNum>
        <w:num w:numId="1">
            <w:abstractNumId w:val="0"/>
        </w:num>
        <w:num w:numId="2">
            <w:abstractNumId w:val="1"/>
        </w:num>
    </w:numbering>
    """

    from lxml import etree
    numbering_elm = etree.fromstring(numbering_xml)

    # Create the numbering part and add to document
    numbering_part = NumberingPart.new()
    numbering_part._element = numbering_elm

    # Add relationship from document to numbering part
    doc.part.relate_to(numbering_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering')


def _apply_list_numbering(para, num_id: int, indent_level: int = 0) -> None:
    """
    Apply Word numbering XML to make paragraph a proper list item.

    Args:
        para: The paragraph to modify
        num_id: The numbering definition ID (1=bullets, 2=numbers)
        indent_level: The indent level for nested lists (0-based)
    """
    # Get or create pPr (paragraph properties)
    pPr = para._p.get_or_add_pPr()

    # Create numPr (numbering properties)
    numPr = OxmlElement('w:numPr')

    # Set indent level
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(indent_level))
    numPr.append(ilvl)

    # Set numId
    numId_elem = OxmlElement('w:numId')
    numId_elem.set(qn('w:val'), str(num_id))
    numPr.append(numId_elem)

    pPr.append(numPr)


def _write_table(doc: DocxDocument, table: Table) -> None:
    """Write a table element."""
    if not table.rows:
        return

    # Create table with appropriate dimensions
    num_rows = len(table.rows)
    num_cols = max(len(row.cells) for row in table.rows) if table.rows else 0

    if num_cols == 0:
        return

    docx_table = doc.add_table(rows=num_rows, cols=num_cols)
    docx_table.style = 'Table Grid'

    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            docx_cell = docx_table.cell(row_idx, col_idx)

            # Clear default paragraph and add our content
            if docx_cell.paragraphs:
                para = docx_cell.paragraphs[0]
                para.clear()
            else:
                para = docx_cell.add_paragraph()

            _add_runs_to_paragraph(para, cell.runs())

            # Apply preserved paragraph formatting
            _apply_paragraph_formatting(para, cell)


def _add_runs_to_paragraph(para, runs: list[TextRun]) -> None:
    """Add formatted TextRuns to a paragraph."""
    for run in runs:
        if run.hyperlink:
            # Hyperlinks need special handling
            _add_hyperlink(para, run)
        else:
            _add_regular_run(para, run)


def _apply_paragraph_formatting(para, element) -> None:
    """
    Apply preserved paragraph-level formatting to a docx paragraph.

    Works for Paragraph, Heading, or ListItem elements that have
    formatting preservation fields.
    """
    pf = para.paragraph_format

    if hasattr(element, 'space_before_pt') and element.space_before_pt is not None:
        pf.space_before = Pt(element.space_before_pt)

    if hasattr(element, 'space_after_pt') and element.space_after_pt is not None:
        pf.space_after = Pt(element.space_after_pt)

    if hasattr(element, 'line_spacing') and element.line_spacing is not None:
        pf.line_spacing = element.line_spacing


def _add_regular_run(para, text_run: TextRun) -> None:
    """Add a regular (non-hyperlink) run with formatting."""
    docx_run = para.add_run(text_run.text)

    # Apply formatting
    if text_run.bold:
        docx_run.bold = True
    if text_run.italic:
        docx_run.italic = True
    if text_run.underline:
        docx_run.underline = True
    if text_run.strikethrough:
        docx_run.font.strike = True

    # Apply highlight
    if text_run.highlight_color:
        color_name = text_run.highlight_color.lower()
        highlight = HIGHLIGHT_MAP.get(color_name)
        if highlight:
            docx_run.font.highlight_color = highlight
        else:
            # Default to yellow if unknown color
            docx_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Apply preserved font formatting
    if text_run.font_name:
        docx_run.font.name = text_run.font_name
    if text_run.font_size_pt is not None:
        docx_run.font.size = Pt(text_run.font_size_pt)


def _add_hyperlink(para, text_run: TextRun) -> None:
    """
    Add a hyperlink run to a paragraph.

    Hyperlinks in OOXML require direct XML manipulation.
    """
    # Get the paragraph's part to access relationships
    part = para.part

    # Create relationship for the hyperlink
    r_id = part.relate_to(
        text_run.hyperlink,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )

    # Create hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create run element inside hyperlink
    new_run = OxmlElement('w:r')

    # Add run properties
    rPr = OxmlElement('w:rPr')

    # Hyperlink styling (blue, underlined)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')  # Blue
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    # Add other formatting
    if text_run.bold:
        bold = OxmlElement('w:b')
        rPr.append(bold)

    if text_run.italic:
        italic = OxmlElement('w:i')
        rPr.append(italic)

    if text_run.strikethrough:
        strike = OxmlElement('w:strike')
        rPr.append(strike)

    if text_run.highlight_color:
        highlight = OxmlElement('w:highlight')
        color_name = text_run.highlight_color.lower()
        # Word uses specific color names
        hl_val = _get_word_highlight_value(color_name)
        highlight.set(qn('w:val'), hl_val)
        rPr.append(highlight)

    new_run.append(rPr)

    # Add text element
    text_elem = OxmlElement('w:t')
    text_elem.text = text_run.text

    # Preserve whitespace
    text_elem.set(qn('xml:space'), 'preserve')

    new_run.append(text_elem)
    hyperlink.append(new_run)

    # Append hyperlink to paragraph
    para._p.append(hyperlink)


def _get_word_highlight_value(color_name: str) -> str:
    """Map color name to Word highlight value string."""
    mapping = {
        'yellow': 'yellow',
        'bright green': 'green',
        'green': 'green',
        'turquoise': 'cyan',
        'cyan': 'cyan',
        'pink': 'magenta',
        'magenta': 'magenta',
        'blue': 'blue',
        'red': 'red',
        'dark blue': 'darkBlue',
        'teal': 'darkCyan',
        'dark red': 'darkRed',
        'dark yellow': 'darkYellow',
        'gray': 'lightGray',
        'gray 25': 'lightGray',
        'gray 50': 'darkGray',
    }
    return mapping.get(color_name, 'yellow')


def write_docx_simple(text: str, output_path: Path) -> None:
    """
    Write plain text to a .docx file.

    Simple utility for testing or when full Document structure isn't needed.
    Each line becomes a paragraph.
    """
    doc = DocxDocument()

    for line in text.split('\n'):
        if line.strip():
            doc.add_paragraph(line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
