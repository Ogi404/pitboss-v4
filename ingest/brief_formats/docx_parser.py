"""
Pitboss v4 - Word Brief Parser

Parses .docx brief files into RawBriefExtraction with confidence scoring.

Handles format variants:
- Inline keywords: "Main: keyword1, keyword2 (3x each)"
- Key-value tables: Two-column tables with field | value
- Structured tables: Multi-column tables with headers
"""

from __future__ import annotations
from pathlib import Path
from typing import Union, Optional, Any
import re

from docx import Document as DocxDocument
from docx.table import Table

from ingest.brief_base import (
    BriefParser,
    register_brief_parser,
    RawBriefExtraction,
    RawKeywordGroup,
    RawSection,
    RawLink,
    is_metadata_label,
    _parse_quantity_range,
    clean_keyword,
)


# Keyword patterns for inline detection
INLINE_KEYWORD_PATTERNS = [
    # "Main keywords: word1, word2, word3"
    re.compile(r"(main|primary|focus)\s*keywords?\s*[:=]\s*(.+)", re.IGNORECASE),
    # "Support keywords: word1, word2"
    re.compile(r"(support|secondary|additional)\s*keywords?\s*[:=]\s*(.+)", re.IGNORECASE),
    # "LSI keywords: word1, word2"
    re.compile(r"(lsi|latent|semantic|related)\s*keywords?\s*[:=]\s*(.+)", re.IGNORECASE),
    # Generic "Keywords: word1, word2"
    re.compile(r"keywords?\s*[:=]\s*(.+)", re.IGNORECASE),
]

# Meta field patterns
META_PATTERNS = {
    "brand": re.compile(r"brand|client|website|site\s*name", re.IGNORECASE),
    "locale": re.compile(r"locale|language|region", re.IGNORECASE),
    "market": re.compile(r"market|country|geo|target", re.IGNORECASE),
    "word_count": re.compile(r"(total\s*)?(word\s*count|words?|length|target\s*length)", re.IGNORECASE),
    "task": re.compile(r"task|topic|article\s*type|content\s*type|type", re.IGNORECASE),
}

# Table header patterns for keywords
KEYWORD_TABLE_HEADERS = [
    (re.compile(r"keyword", re.IGNORECASE), re.compile(r"(qty|quantity|count|#)", re.IGNORECASE)),
    (re.compile(r"main", re.IGNORECASE), re.compile(r"(qty|quantity|count)", re.IGNORECASE)),
]

# Table header patterns for sections
SECTION_TABLE_HEADERS = [
    (re.compile(r"section|heading|structure", re.IGNORECASE), re.compile(r"word\s*count|words?|length", re.IGNORECASE)),
]

# Group name patterns
GROUP_PATTERNS = {
    "main": re.compile(r"\bmain\b|primary|focus", re.IGNORECASE),
    "support": re.compile(r"support|secondary|additional", re.IGNORECASE),
    "lsi": re.compile(r"\blsi\b|latent|semantic|related", re.IGNORECASE),
}


def _parse_quantity(value: Any) -> int:
    """Parse quantity from value."""
    if value is None:
        return 1
    if isinstance(value, (int, float)):
        return max(1, int(value))
    text = str(value).strip()
    match = re.search(r"(\d+)", text)
    if match:
        return max(1, int(match.group(1)))
    return 1


def _infer_group_name(header: str) -> str:
    """Infer keyword group name from header text."""
    for group_name, pattern in GROUP_PATTERNS.items():
        if pattern.search(header):
            return group_name
    return "main"


def _strip_translations(text: str) -> str:
    """
    Strip translation blocks from multi-language keyword lines.

    Format: "eng1, eng2, eng3. (ITA: ita1, ita2.) czech1, czech2 (CZ)"
    Also handles: "eng1, eng2.ITA: ita1, ita2" (no space before marker)

    Returns only the English portion before the first parenthetical or language marker.
    """
    # Find the first opening parenthesis - everything after is translations
    paren_idx = text.find("(")
    if paren_idx > 0:
        text = text[:paren_idx]

    # Also check for language markers without parentheses (e.g., ".ITA:" or " ITA:")
    lang_markers = ["ITA:", "CZ:", "PT:", "ES:", "DE:", "FR:", "(CZ)", "(ITA)"]
    for marker in lang_markers:
        marker_idx = text.upper().find(marker.upper())
        if marker_idx > 0:
            text = text[:marker_idx]

    # Strip trailing period and whitespace
    return text.rstrip(". \t")


def _is_valid_keyword(kw: str) -> bool:
    """
    Validate that a keyword is clean and not a translation fragment.

    Returns False if:
    - Contains language markers (ITA:, CZ, parentheses)
    - Is longer than 8 words (likely a fragment/sentence)
    - Is empty or metadata
    """
    if not kw:
        return False

    # Check for language markers
    if re.search(r"\b(ITA|CZ|PT|ES|DE|FR)\s*:", kw, re.IGNORECASE):
        return False
    if "(" in kw or ")" in kw:
        return False

    # Check word count - keywords are typically 1-6 words
    word_count = len(kw.split())
    if word_count > 8:
        return False

    # Check for metadata labels
    if is_metadata_label(kw):
        return False

    return True


def _parse_inline_keywords(text: str) -> list[tuple[str, Optional[int], Optional[int]]]:
    """Parse comma/semicolon/slash-separated keywords with optional quantities."""
    keywords = []

    # BUG 3 FIX: Strip translations before parsing
    text = _strip_translations(text)

    # Split on comma, semicolon, OR forward slash
    parts = re.split(r"[,;/]", text)

    for part in parts:
        part = part.strip()
        if not part:
            continue

        # Filter out metadata labels
        if is_metadata_label(part):
            continue

        # Check for quantity patterns
        # "keyword (3x)" or "keyword x3" or "keyword - 3"
        qty_match = re.match(r"(.+?)\s*[\(\-]\s*(\d+)\s*[x×]?\s*\)?$", part)
        if qty_match:
            kw = clean_keyword(qty_match.group(1).strip())
            qty = int(qty_match.group(2))
            if _is_valid_keyword(kw):
                keywords.append((kw, qty, qty))  # Exact quantity
        else:
            qty_match = re.match(r"(.+?)\s*[x×]\s*(\d+)$", part)
            if qty_match:
                kw = clean_keyword(qty_match.group(1).strip())
                qty = int(qty_match.group(2))
                if _is_valid_keyword(kw):
                    keywords.append((kw, qty, qty))  # Exact quantity
            else:
                kw = clean_keyword(part)
                if _is_valid_keyword(kw):
                    keywords.append((kw, None, None))  # No quantity constraint

    return keywords


@register_brief_parser
class DocxBriefParser(BriefParser):
    """Parser for Word (.docx) brief files."""

    def get_format_name(self) -> str:
        return "docx"

    def can_parse(self, source: Union[Path, str]) -> bool:
        if isinstance(source, str):
            source = Path(source)
        return source.suffix.lower() == ".docx"

    def extract(self, source: Union[Path, str]) -> RawBriefExtraction:
        """Extract raw data from Word brief."""
        if isinstance(source, str):
            source = Path(source)

        doc = DocxDocument(str(source))

        extraction = RawBriefExtraction(
            source_path=str(source),
            source_format="docx",
        )

        # BUG 1 FIX: Detect multi-task briefs via Title-styled paragraphs
        tasks = self._detect_tasks_by_title_style(doc)
        if tasks:
            extraction.tasks = tasks

        # Extract from tables first (higher confidence)
        table_keywords, table_kw_conf = self._extract_keywords_from_tables(doc)
        table_sections, table_sec_conf = self._extract_sections_from_tables(doc)
        table_meta = self._extract_meta_from_tables(doc)

        # Extract from prose (lower confidence fallback)
        prose_keywords, prose_kw_conf = self._extract_keywords_from_prose(doc)
        prose_meta = self._extract_meta_from_prose(doc)

        # Combine results, preferring table data
        if table_keywords:
            extraction.keyword_groups = table_keywords
            extraction.keywords_confidence = table_kw_conf
        elif prose_keywords:
            extraction.keyword_groups = prose_keywords
            extraction.keywords_confidence = prose_kw_conf

        if table_sections:
            extraction.sections = table_sections
            extraction.sections_confidence = table_sec_conf

        # Merge meta (table takes precedence)
        meta = {**prose_meta, **table_meta}
        extraction.brand_name = meta.get("brand")
        extraction.brand_confidence = meta.get("brand_confidence", 0.5)
        extraction.locale = meta.get("locale")
        extraction.market = meta.get("market")
        extraction.locale_confidence = meta.get("locale_confidence", 0.5)
        extraction.target_word_count = meta.get("word_count")
        extraction.word_count_confidence = meta.get("word_count_confidence", 0.5)
        extraction.task_name = meta.get("task")
        extraction.task_name_confidence = meta.get("task_confidence", 0.5)

        # Store raw for debugging
        extraction.raw_data = {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
        }

        return extraction

    def _extract_keywords_from_tables(self, doc: DocxDocument) -> tuple[list[RawKeywordGroup], float]:
        """Extract keywords from table format."""
        groups = []
        confidence = 0.0

        for table in doc.tables:
            table_groups, table_conf = self._parse_keyword_table(table)
            if table_groups:
                groups.extend(table_groups)
                confidence = max(confidence, table_conf)

        return groups, confidence

    def _parse_keyword_table(self, table: Table) -> tuple[list[RawKeywordGroup], float]:
        """Parse a table for keyword data."""
        groups = []
        confidence = 0.0

        if len(table.rows) < 2:
            return groups, confidence

        # Check first row for headers
        header_row = table.rows[0]
        header_cells = [cell.text.strip() for cell in header_row.cells]

        # Look for keyword/quantity column pairs
        for kw_pattern, qty_pattern in KEYWORD_TABLE_HEADERS:
            kw_col = None
            qty_col = None

            for idx, header in enumerate(header_cells):
                if kw_col is None and kw_pattern.search(header):
                    kw_col = idx
                if qty_col is None and qty_pattern.search(header):
                    qty_col = idx

            if kw_col is not None:
                # Extract keywords from data rows
                keywords = []
                for row in table.rows[1:]:
                    cells = row.cells
                    if kw_col < len(cells):
                        kw_text = cells[kw_col].text.strip()
                        if kw_text and not is_metadata_label(kw_text):
                            kw_clean = clean_keyword(kw_text)
                            if kw_clean and not is_metadata_label(kw_clean):
                                min_qty, max_qty = None, None
                                if qty_col is not None and qty_col < len(cells):
                                    min_qty, max_qty = _parse_quantity_range(cells[qty_col].text)
                                keywords.append((kw_clean, min_qty, max_qty))

                if keywords:
                    group_name = _infer_group_name(header_cells[kw_col])
                    groups.append(RawKeywordGroup(
                        keywords=keywords,
                        group_name=group_name,
                        confidence=0.90 if qty_col else 0.75,
                    ))
                    confidence = max(confidence, 0.90 if qty_col else 0.75)

        return groups, confidence

    def _extract_sections_from_tables(self, doc: DocxDocument) -> tuple[list[RawSection], float]:
        """Extract sections from table format."""
        sections = []
        confidence = 0.0

        for table in doc.tables:
            table_sections, table_conf = self._parse_section_table(table)
            if table_sections:
                sections.extend(table_sections)
                confidence = max(confidence, table_conf)

        return sections, confidence

    def _parse_section_table(self, table: Table) -> tuple[list[RawSection], float]:
        """Parse a table for section structure."""
        sections = []
        confidence = 0.0

        if len(table.rows) < 2:
            return sections, confidence

        # Check first row for headers
        header_row = table.rows[0]
        header_cells = [cell.text.strip() for cell in header_row.cells]

        # Look for section/word count column pairs
        for sec_pattern, wc_pattern in SECTION_TABLE_HEADERS:
            sec_col = None
            wc_col = None

            for idx, header in enumerate(header_cells):
                if sec_col is None and sec_pattern.search(header):
                    sec_col = idx
                if wc_col is None and wc_pattern.search(header):
                    wc_col = idx

            if sec_col is not None:
                # Extract sections from data rows
                for row in table.rows[1:]:
                    cells = row.cells
                    if sec_col < len(cells):
                        sec_text = cells[sec_col].text.strip()
                        if sec_text:
                            word_count = None
                            if wc_col is not None and wc_col < len(cells):
                                word_count = _parse_quantity(cells[wc_col].text)
                                if word_count == 1:  # Default, likely not a word count
                                    word_count = None

                            sections.append(RawSection(
                                heading=sec_text,
                                word_count=word_count,
                                confidence=0.90 if wc_col and word_count else 0.70,
                            ))
                            confidence = max(confidence, 0.90 if wc_col and word_count else 0.70)

        return sections, confidence

    def _extract_meta_from_tables(self, doc: DocxDocument) -> dict[str, Any]:
        """Extract meta fields from key-value tables."""
        meta = {}

        for table in doc.tables:
            # Key-value tables typically have 2 columns
            if len(table.columns) == 2:
                for row in table.rows:
                    cells = row.cells
                    if len(cells) >= 2:
                        key = cells[0].text.strip()
                        value = cells[1].text.strip()

                        if not key or not value:
                            continue

                        for field_name, pattern in META_PATTERNS.items():
                            if pattern.search(key) and field_name not in meta:
                                if field_name == "word_count":
                                    meta[field_name] = _parse_quantity(value)
                                    meta[f"{field_name}_confidence"] = 0.90
                                else:
                                    meta[field_name] = value
                                    meta[f"{field_name}_confidence"] = 0.85

        return meta

    def _extract_keywords_from_prose(self, doc: DocxDocument) -> tuple[list[RawKeywordGroup], float]:
        """Extract keywords from prose paragraphs."""
        groups = []
        confidence = 0.0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            for pattern in INLINE_KEYWORD_PATTERNS:
                match = pattern.search(text)
                if match:
                    # Determine group name from pattern match
                    if len(match.groups()) >= 2:
                        group_hint = match.group(1)
                        keyword_text = match.group(2)
                    else:
                        group_hint = "main"
                        keyword_text = match.group(1)

                    keywords = _parse_inline_keywords(keyword_text)
                    if keywords:
                        group_name = _infer_group_name(group_hint)
                        groups.append(RawKeywordGroup(
                            keywords=keywords,
                            group_name=group_name,
                            confidence=0.60,  # Lower confidence for inline
                        ))
                        confidence = max(confidence, 0.60)
                    break  # Only match one pattern per paragraph

        return groups, confidence

    def _extract_meta_from_prose(self, doc: DocxDocument) -> dict[str, Any]:
        """Extract meta fields from prose paragraphs."""
        meta = {}

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Look for "Field: Value" patterns
            for field_name, pattern in META_PATTERNS.items():
                if field_name in meta:
                    continue

                # Try "Field: Value" pattern - use non-capturing wrapper for the field pattern
                inline_match = re.match(rf"(?:{pattern.pattern})\s*[:=]\s*(.+)", text, re.IGNORECASE)
                if inline_match:
                    # The value is in the last group
                    value = inline_match.group(inline_match.lastindex)
                    if value:
                        value = value.strip()
                        if field_name == "word_count":
                            meta[field_name] = _parse_quantity(value)
                            meta[f"{field_name}_confidence"] = 0.70
                        else:
                            meta[field_name] = value
                            meta[f"{field_name}_confidence"] = 0.65

        return meta

    def _detect_tasks_by_title_style(self, doc: DocxDocument) -> list[str]:
        """
        BUG 1 FIX: Detect multi-task briefs by Title-styled paragraphs.

        Multi-page docx briefs like "Content task Mar'26 22bet.com.gh - 10 pages.docx"
        use Title-styled paragraphs to delimit each task/page block.

        Returns list of task names if multiple Title paragraphs found, else empty list.
        """
        tasks = []

        for para in doc.paragraphs:
            # Check paragraph style
            style_name = para.style.name if para.style else ""

            # Title style marks the start of a new task block
            if style_name == "Title":
                task_name = para.text.strip()
                if task_name and task_name not in tasks:
                    # Filter out generic titles that aren't task names
                    # (e.g., document titles like "Content Brief")
                    if not re.match(r"^(content\s+)?(brief|task|plan)s?$", task_name, re.IGNORECASE):
                        tasks.append(task_name)

        # Only return if we have multiple tasks (single task = not multi-task)
        return tasks if len(tasks) > 1 else []
