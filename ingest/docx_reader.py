"""
Pitboss v4 - DOCX Reader

Parses .docx files into Document objects for corpus analysis.
Extracts paragraphs, headings, lists, tables, and inline formatting.
"""

from __future__ import annotations
from pathlib import Path
from typing import Optional
import re

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from core.document import (
    Document,
    Paragraph,
    Heading,
    HeadingLevel,
    List,
    ListItem,
    ListType,
    Table,
    TableRow,
    TableCell,
    TextRun,
    BlockElement,
)


# Heading style name patterns
HEADING_PATTERNS = {
    1: re.compile(r'heading\s*1|title', re.IGNORECASE),
    2: re.compile(r'heading\s*2', re.IGNORECASE),
    3: re.compile(r'heading\s*3', re.IGNORECASE),
    4: re.compile(r'heading\s*4', re.IGNORECASE),
}


def _detect_heading_level(style_name: Optional[str]) -> Optional[HeadingLevel]:
    """Detect heading level from Word style name."""
    if not style_name:
        return None

    for level, pattern in HEADING_PATTERNS.items():
        if pattern.search(style_name):
            return HeadingLevel(level)

    return None


def _get_hyperlink_url(run) -> Optional[str]:
    """Extract hyperlink URL from a run if present."""
    # Check if run is inside a hyperlink
    parent = run._r.getparent()
    if parent is not None:
        # Look for w:hyperlink ancestor
        hyperlink = parent.getparent()
        if hyperlink is not None and hyperlink.tag.endswith('hyperlink'):
            # Get the relationship ID
            r_id = hyperlink.get(qn('r:id'))
            if r_id:
                # Get the actual URL from relationships
                try:
                    part = run.part
                    rel = part.rels.get(r_id)
                    if rel and hasattr(rel, 'target_ref'):
                        return rel.target_ref
                except Exception:
                    pass
    return None


def _extract_highlight_color(run) -> Optional[str]:
    """Extract highlight color from a run."""
    try:
        highlight = run.font.highlight_color
        if highlight:
            return str(highlight).lower().replace('_', ' ')
    except Exception:
        pass
    return None


def _extract_runs(paragraph, start_offset: int = 0) -> list[TextRun]:
    """Extract TextRun objects from a docx paragraph."""
    runs = []
    current_offset = 0

    for run in paragraph.runs:
        text = run.text
        if not text:
            continue

        # Extract formatting
        bold = run.bold or False
        italic = run.italic or False
        underline = run.underline is not None and run.underline
        strikethrough = run.font.strike or False
        highlight = _extract_highlight_color(run)
        hyperlink = _get_hyperlink_url(run)

        runs.append(TextRun(
            text=text,
            start_offset=current_offset,
            end_offset=current_offset + len(text),
            bold=bold,
            italic=italic,
            underline=underline,
            strikethrough=strikethrough,
            highlight_color=highlight,
            hyperlink=hyperlink,
        ))

        current_offset += len(text)

    return runs


def _is_list_paragraph(paragraph) -> tuple[bool, Optional[ListType], int]:
    """
    Check if a paragraph is a list item.

    Returns (is_list, list_type, indent_level)
    """
    # Check for numbering in paragraph XML
    p = paragraph._p
    numPr = p.find(qn('w:pPr'))
    if numPr is not None:
        numPr_elem = numPr.find(qn('w:numPr'))
        if numPr_elem is not None:
            ilvl = numPr_elem.find(qn('w:ilvl'))
            indent_level = int(ilvl.get(qn('w:val'))) if ilvl is not None else 0

            # Try to determine if ordered or unordered
            # This is simplified - full detection requires parsing numbering.xml
            numId = numPr_elem.find(qn('w:numId'))
            if numId is not None:
                # Default to unordered, but could be ordered
                return True, ListType.UNORDERED, indent_level

    # Check for bullet characters at start of text
    text = paragraph.text.strip()
    if text.startswith(('•', '-', '*', '○', '●')):
        return True, ListType.UNORDERED, 0
    if re.match(r'^\d+[\.\)]\s', text):
        return True, ListType.ORDERED, 0

    return False, None, 0


def read_docx(filepath: Path) -> Document:
    """
    Parse a .docx file into a Document object.

    Extracts:
    - Paragraphs with text and character offsets
    - Headings (H1-H4) detected by Word styles
    - Lists (ordered/unordered)
    - Tables with cells
    - Basic formatting runs (bold, italic, hyperlinks, highlights)

    Args:
        filepath: Path to the .docx file

    Returns:
        Document with all elements in document order
    """
    doc = DocxDocument(str(filepath))
    elements: list[BlockElement] = []
    current_offset = 0

    # Track list accumulation
    pending_list_items: list[ListItem] = []
    pending_list_type: Optional[ListType] = None
    list_start_offset: int = 0

    def flush_list():
        """Flush accumulated list items to elements."""
        nonlocal pending_list_items, pending_list_type, list_start_offset
        if pending_list_items:
            list_end = pending_list_items[-1].end_offset
            elements.append(List(
                list_type=pending_list_type or ListType.UNORDERED,
                items=pending_list_items,
                start_offset=list_start_offset,
                end_offset=list_end,
            ))
            pending_list_items = []
            pending_list_type = None

    # Process paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            # Empty paragraph - flush any pending list
            flush_list()
            continue

        # Check for heading
        style_name = para.style.name if para.style else None
        heading_level = _detect_heading_level(style_name)

        # Check for list
        is_list, list_type, indent_level = _is_list_paragraph(para)

        # Calculate offsets
        text_len = len(text)
        start = current_offset
        end = current_offset + text_len

        # Extract formatting runs
        runs = _extract_runs(para, start)

        if heading_level:
            # Flush any pending list before heading
            flush_list()
            elements.append(Heading(
                text=text,
                level=heading_level,
                start_offset=start,
                end_offset=end,
                _runs=runs,
            ))
        elif is_list:
            # Accumulate list items
            if not pending_list_items:
                list_start_offset = start
                pending_list_type = list_type

            pending_list_items.append(ListItem(
                text=text,
                start_offset=start,
                end_offset=end,
                indent_level=indent_level,
                _runs=runs,
            ))
        else:
            # Regular paragraph
            flush_list()
            elements.append(Paragraph(
                text=text,
                start_offset=start,
                end_offset=end,
                _runs=runs,
            ))

        current_offset = end + 1  # +1 for newline separator

    # Flush any remaining list
    flush_list()

    # Process tables
    for table in doc.tables:
        table_start = current_offset
        rows = []

        for row_idx, row in enumerate(table.rows):
            cells = []
            is_header_row = row_idx == 0  # First row assumed to be header

            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                cell_start = current_offset
                cell_end = current_offset + len(cell_text)

                # Extract runs from first paragraph of cell
                cell_runs = []
                if cell.paragraphs:
                    cell_runs = _extract_runs(cell.paragraphs[0], 0)

                cells.append(TableCell(
                    text=cell_text,
                    start_offset=cell_start,
                    end_offset=cell_end,
                    row_index=row_idx,
                    col_index=col_idx,
                    is_header=is_header_row,
                    _runs=cell_runs,
                ))

                current_offset = cell_end + 1

            rows.append(TableRow(
                cells=cells,
                is_header_row=is_header_row,
            ))

        table_end = current_offset
        elements.append(Table(
            rows=rows,
            start_offset=table_start,
            end_offset=table_end,
        ))

    return Document.from_elements(
        elements=elements,
        title=filepath.stem,
        source_url=str(filepath),
        source_format="docx",
    )


def read_docx_text_only(filepath: Path) -> str:
    """
    Read a .docx file and return just the plain text.

    Simpler alternative when full Document parsing is not needed.
    """
    doc = DocxDocument(str(filepath))
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    return "\n".join(paragraphs)


def get_docx_headings(filepath: Path) -> list[tuple[str, HeadingLevel]]:
    """
    Extract just the headings from a .docx file.

    Useful for quick structure analysis.

    Returns list of (heading_text, level) tuples.
    """
    doc = DocxDocument(str(filepath))
    headings = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else None
        level = _detect_heading_level(style_name)

        if level:
            headings.append((text, level))

    return headings
