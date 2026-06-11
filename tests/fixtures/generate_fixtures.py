"""
Generate test fixture .docx files for voice model testing.

Run this script to create/recreate the test fixtures:
    python tests/fixtures/generate_fixtures.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE


def create_docx(filepath: Path, content: list[dict]):
    """
    Create a .docx file with specified content.

    Args:
        filepath: Path to save the file
        content: List of dicts with 'type' and 'text' keys
                 type can be: 'heading1', 'heading2', 'heading3', 'paragraph', 'bullet'
    """
    doc = Document()

    for item in content:
        item_type = item.get('type', 'paragraph')
        text = item.get('text', '')

        if item_type == 'heading1':
            doc.add_heading(text, level=1)
        elif item_type == 'heading2':
            doc.add_heading(text, level=2)
        elif item_type == 'heading3':
            doc.add_heading(text, level=3)
        elif item_type == 'bullet':
            doc.add_paragraph(text, style='List Bullet')
        else:
            doc.add_paragraph(text)

    filepath.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(filepath))
    print(f"Created: {filepath}")


def generate_all_fixtures():
    """Generate all test fixture files."""
    fixtures_dir = Path(__file__).parent

    # =========================================================================
    # BRAND ALPHA (3 articles - below threshold)
    # =========================================================================
    alpha_dir = fixtures_dir / "brand_alpha"

    # Alpha main review - mostly second person
    create_docx(alpha_dir / "alpha-main-review.docx", [
        {"type": "heading1", "text": "Alpha Casino Review"},
        {"type": "paragraph", "text": "Welcome to Alpha Casino. You will find everything you need here. Your experience matters to us."},
        {"type": "heading2", "text": "Getting Started"},
        {"type": "paragraph", "text": "When you sign up, you get access to all games. You can play slots and table games."},
        {"type": "paragraph", "text": "You should check out our promotions. Your first deposit is matched one hundred percent."},
        {"type": "heading2", "text": "Games Available"},
        {"type": "paragraph", "text": "You have many options here. You can choose from slots and poker games. Your favorites are here."},
        {"type": "paragraph", "text": "Additionally, you can try live dealer games. You will enjoy the experience."},
    ])

    # Alpha bonus page
    create_docx(alpha_dir / "alpha-bonus-promo.docx", [
        {"type": "heading1", "text": "Alpha Casino Bonus Offers"},
        {"type": "paragraph", "text": "You can claim your welcome bonus today. Your bonus is waiting for you now."},
        {"type": "heading2", "text": "How to Claim"},
        {"type": "paragraph", "text": "First, you need to register. Then you make your deposit and bonus is yours."},
        {"type": "paragraph", "text": "You must wager thirty five times before withdrawal. You can check terms below."},
    ])

    # Alpha app review
    create_docx(alpha_dir / "alpha-mobile-app.docx", [
        {"type": "heading1", "text": "Alpha Casino Mobile App"},
        {"type": "paragraph", "text": "You can download the app from app store. Your phone will work great with it."},
        {"type": "heading2", "text": "Features"},
        {"type": "paragraph", "text": "You get full access to games. You can deposit and withdraw from mobile. Your account is secure."},
        {"type": "paragraph", "text": "Furthermore, you can enable notifications. You will never miss a promotion."},
    ])

    # =========================================================================
    # BRAND BETA (12 articles - above threshold)
    # Known characteristics: 10-word sentences, title case headings
    # =========================================================================
    beta_dir = fixtures_dir / "brand_beta"

    # Generate 12 articles for beta
    beta_articles = [
        ("beta-main-review.docx", "main_review", "Beta Casino Review"),
        ("beta-bonus.docx", "bonus_page", "Beta Welcome Bonus"),
        ("beta-app.docx", "app_review", "Beta Mobile App"),
        ("beta-slots.docx", "game_review", "Beta Slot Games"),
        ("beta-payments.docx", "payments", "Beta Payment Methods"),
        ("beta-support.docx", "customer_support", "Beta Customer Support"),
        ("beta-vip.docx", "vip_loyalty", "Beta VIP Program"),
        ("beta-boxing.docx", "sports_market", "Beta Boxing Betting"),
        ("beta-registration.docx", "registration", "Beta Sign Up Guide"),
        ("beta-responsible.docx", "responsible_gaming", "Beta Responsible Gaming"),
        ("beta-poker.docx", "game_review", "Beta Poker Games"),
        ("beta-live.docx", "game_review", "Beta Live Casino"),
    ]

    for filename, article_type, title in beta_articles:
        # Create content with exactly 10-word sentences (for testing)
        content = [
            {"type": "heading1", "text": title},
            {"type": "paragraph", "text": "You will find excellent service at Beta Casino today. Your experience here will be amazing and worth your time."},
            {"type": "heading2", "text": "Main Features"},
            {"type": "paragraph", "text": "Beta offers you many great options for your enjoyment. You can access games anytime from anywhere you want."},
            {"type": "paragraph", "text": "Your security is our top priority at Beta Casino. We ensure your data is always protected and safe."},
            {"type": "heading2", "text": "Getting Started Today"},
            {"type": "paragraph", "text": "You need to register first to start playing here. The process takes only a few minutes to complete."},
        ]
        create_docx(beta_dir / filename, content)

    # =========================================================================
    # BRAND GAMMA (5 articles - below threshold, all sports_market)
    # =========================================================================
    gamma_dir = fixtures_dir / "brand_gamma"

    sports = ["boxing", "basketball", "football", "tennis", "cricket"]
    for sport in sports:
        create_docx(gamma_dir / f"gamma-{sport}-betting.docx", [
            {"type": "heading1", "text": f"Gamma {sport.title()} Betting Guide"},
            {"type": "paragraph", "text": f"You can bet on {sport} matches at Gamma. Your bets are processed instantly here."},
            {"type": "heading2", "text": "Available Markets"},
            {"type": "paragraph", "text": f"You have many {sport} betting options available. You can bet on match winners and more."},
            {"type": "paragraph", "text": "However, you should check the odds first. You will find competitive prices here."},
            {"type": "heading2", "text": "Live Betting"},
            {"type": "paragraph", "text": f"You can place live bets during {sport} events. Your bets are settled quickly here."},
        ])

    print("\n=== Fixture Generation Complete ===")
    print(f"Brand Alpha: 3 articles (below threshold)")
    print(f"Brand Beta: 12 articles (above threshold)")
    print(f"Brand Gamma: 5 articles (below threshold, all sports_market)")


if __name__ == "__main__":
    generate_all_fixtures()
